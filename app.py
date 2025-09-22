import streamlit as st
from io import BytesIO
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from PIL import Image as PILImage
import tempfile, base64, zipfile

st.set_page_config(page_title="Resume Builder", layout="wide")
st.markdown("<h1 style='text-align:center;color:#2E86C1;'>📄 Professional Resume Builder</h1>", unsafe_allow_html=True)
with st.sidebar:
    st.header("👤 Profile Information")
    name = st.text_input("Full Name", "Jane Doe")
    title = st.text_input("Job Title", "Software Engineer")
    email = st.text_input("Email", "jane@example.com")
    phone = st.text_input("Phone", "+1-234-567-890")
    linkedin = st.text_input("LinkedIn", "https://linkedin.com/in/janedoe")
    github = st.text_input("GitHub", "https://github.com/janedoe")
    photo = st.file_uploader("Upload Profile Photo", type=["png","jpg","jpeg"])
    st.header("📑 Resume Content")
    summary = st.text_area("Professional Summary", "Results-driven engineer with proven experience...")
    education = st.text_area("Education", "B.Sc. Computer Science — XYZ University (2018-2022)")
    experience = st.text_area("Experience", "Software Engineer — ACME Inc. (2022-Present)")
    projects = st.text_area("Projects", "- Project A\n- Project B")
    activities = st.text_area("Activities", "- Member of ABC Club")
    skills = st.text_area("Skills", "Python, SQL, React, Leadership")
    certifications = st.text_area("Certifications", "AWS Certified Developer")
    languages = st.text_area("Languages", "English, Spanish")
    st.header("🎨 Layout & Style")
    layout = st.radio("Layout", ["Single Column", "Two Column"], index=0)
    margin_choice = st.radio("Page Margin", ["Narrow", "Normal", "Wide"], index=1)
    template = st.selectbox("Resume Template", ["Classic", "Modern", "Minimalist"])
    accent_color = st.color_picker("Accent Color", "#2E86C1")
    generate = st.button("🚀 Generate Resume")

def get_margins(choice):
    return dict(left=30,right=30,top=40,bottom=40) if choice=="Narrow" else dict(left=80,right=80,top=80,bottom=80) if choice=="Wide" else dict(left=50,right=50,top=50,bottom=50)

def bullet_list(items, style):
    return ListFlowable([ListItem(Paragraph(i.strip(), style)) for i in items.split("\n") if i.strip()], bulletType='bullet')

def create_pdf(data, photo_file, layout, template, margins, accent_color):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=margins["left"], rightMargin=margins["right"], topMargin=margins["top"], bottomMargin=margins["bottom"])
    styles = getSampleStyleSheet()
    normal = ParagraphStyle('normal', parent=styles['Normal'], fontSize=10, leading=12)
    heading = ParagraphStyle('heading', parent=styles['Heading2'], fontSize=12, leading=14, spaceAfter=6, textColor=colors.HexColor(accent_color), underlineWidth=0.5 if template=="Modern" else 0)
    story = []
    if photo_file:
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        PILImage.open(photo_file).save(tmp.name)
        im = Image(tmp.name, width=1.0*inch, height=1.0*inch)
        header_table = Table([[im, Paragraph(f"<b>{data['name']}</b><br/>{data['title']}<br/>{data['email']} | {data['phone']}<br/><a href='{data['linkedin']}'>LinkedIn</a> | <a href='{data['github']}'>GitHub</a>", normal)]], colWidths=[1.2*inch, 5*inch])
        header_table.setStyle(TableStyle([("VALIGN",(0,0),(-1,-1),"TOP")]))
        story.append(header_table)
    else:
        story.append(Paragraph(f"<b>{data['name']}</b><br/>{data['title']}<br/>{data['email']} | {data['phone']}<br/><a href='{data['linkedin']}'>LinkedIn</a> | <a href='{data['github']}'>GitHub</a>", normal))
    story.append(Spacer(1,12))
    def section(title, content, bullets=False):
        story.append(Paragraph(title, heading))
        if bullets: story.append(bullet_list(content, normal))
        else: story.append(Paragraph(content.replace("\n","<br/>"), normal))
        story.append(Spacer(1,8))
    if layout=="Single Column":
        section("Professional Summary", data['summary'])
        section("Education", data['education'])
        section("Experience", data['experience'])
        section("Projects", data['projects'], bullets=True)
        section("Activities", data['activities'], bullets=True)
        section("Skills", data['skills'], bullets=True)
        section("Certifications", data['certifications'], bullets=True)
        section("Languages", data['languages'], bullets=True)
    else:
        left, right = [], []
        left.append(Paragraph("Summary", heading)); left.append(Paragraph(data['summary'], normal)); left.append(Spacer(1,6))
        left.append(Paragraph("Skills", heading)); left.append(bullet_list(data['skills'], normal))
        left.append(Paragraph("Certifications", heading)); left.append(bullet_list(data['certifications'], normal))
        left.append(Paragraph("Languages", heading)); left.append(bullet_list(data['languages'], normal))
        right.append(Paragraph("Education", heading)); right.append(Paragraph(data['education'], normal))
        right.append(Paragraph("Experience", heading)); right.append(Paragraph(data['experience'], normal))
        right.append(Paragraph("Projects", heading)); right.append(bullet_list(data['projects'], normal))
        right.append(Paragraph("Activities", heading)); right.append(bullet_list(data['activities'], normal))
        table = Table([[left, right]], colWidths=[2.5*inch, 3.5*inch])
        table.setStyle(TableStyle([("VALIGN",(0,0),(-1,-1),"TOP"), ("LEFTPADDING",(0,0),(-1,-1),6), ("RIGHTPADDING",(0,0),(-1,-1),6)]))
        story.append(table)
    doc.build(story); buffer.seek(0); return buffer.getvalue()

def create_docx(data, photo_file):
    doc = Document()
    if photo_file:
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        PILImage.open(photo_file).save(tmp.name)
        doc.add_picture(tmp.name, width=1.0*inch)
    doc.add_heading(data['name'], 0); doc.add_paragraph(data['title']); doc.add_paragraph(f"{data['email']} | {data['phone']}\n{data['linkedin']} | {data['github']}")
    def section(title, content, bullets=False):
        doc.add_heading(title, level=1)
        if bullets: [doc.add_paragraph(i.strip(), style="List Bullet") for i in content.split("\n") if i.strip()]
        else: doc.add_paragraph(content)
    section("Summary", data['summary'])
    section("Education", data['education'])
    section("Experience", data['experience'])
    section("Projects", data['projects'], bullets=True)
    section("Activities", data['activities'], bullets=True)
    section("Skills", data['skills'], bullets=True)
    section("Certifications", data['certifications'], bullets=True)
    section("Languages", data['languages'], bullets=True)
    buf = BytesIO(); doc.save(buf); buf.seek(0); return buf.getvalue()

if generate:
    data = {"name":name,"title":title,"email":email,"phone":phone,"linkedin":linkedin,"github":github,"summary":summary,"education":education,"experience":experience,"projects":projects,"activities":activities,"skills":skills,"certifications":certifications,"languages":languages}
    pdf_bytes = create_pdf(data, photo, layout, template, get_margins(margin_choice), accent_color)
    docx_bytes = create_docx(data, photo)
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer,"w") as zf:
        zf.writestr("resume.pdf", pdf_bytes)
        zf.writestr("resume.docx", docx_bytes)
    zip_buffer.seek(0)
    st.download_button("📥 Download PDF", data=pdf_bytes, file_name="resume.pdf", mime="application/pdf")
    st.download_button("📥 Download DOCX", data=docx_bytes, file_name="resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    st.download_button("📦 Download All (ZIP)", data=zip_buffer, file_name="resume.zip", mime="application/zip")
    st.success("✅ Resume generated successfully!")
